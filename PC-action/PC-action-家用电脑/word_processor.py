import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


class WordProcessor:
    """Word文档处理类，提供清除表格边框、设置字体和清理内容等功能"""
    
    def __init__(self, file_path):
        """
        初始化Word处理器
        :param file_path: Word文档路径
        """
        self.file_path = file_path
        self.doc = Document(file_path)
    
    def clear_table_borders(self):
        """清除所有表格的内外边框"""
        # 首先清除所有段落的边框
        self._clear_paragraph_borders()
        
        # 然后处理表格边框
        for table in self.doc.tables:
            # 清除表格边框
            tbl = table._tbl
            tblPr = tbl.tblPr
            
            # 如果没有tblPr，创建一个
            if tblPr is None:
                from docx.oxml import parse_xml
                tblPr = parse_xml('<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                tbl.insert(0, tblPr)
            
            # 直接移除所有tblBorders元素
            tblBorders_list = tblPr.xpath(".//w:tblBorders")
            for tblBorders in tblBorders_list:
                tblPr.remove(tblBorders)
            
            # 清除每个单元格的边框
            for row in table.rows:
                for cell in row.cells:
                    self._clear_cell_borders(cell)
            
            # 尝试清除表格样式中的边框设置
            self._clear_table_style_borders(table)
            
            # 清除表格中的底纹，这可能会被误认为是边框
            self._clear_table_shading(table)
    
    def _clear_table_shading(self, table):
        """清除表格中的底纹，这可能会被误认为是边框"""
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.tcPr
                
                if tcPr is None:
                    continue
                
                # 移除所有shading元素
                shading_list = tcPr.xpath(".//w:shading")
                for shading in shading_list:
                    tcPr.remove(shading)
                
                # 也检查单元格内段落的底纹
                for paragraph in cell.paragraphs:
                    self._clear_paragraph_shading(paragraph)
    
    def _clear_paragraph_shading(self, paragraph):
        """清除段落的底纹"""
        pPr = paragraph._p.pPr
        if pPr is None:
            return
        
        # 移除所有shading元素
        shading_list = pPr.xpath(".//w:shading")
        for shading in shading_list:
            pPr.remove(shading)
    
    def _clear_table_style_borders(self, table):
        """清除表格样式中的边框设置"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        
        if tblPr is None:
            return
        
        # 查找并移除表格样式引用
        tblStyle = tblPr.first_child_found_in("w:tblStyle")
        if tblStyle is not None:
            # 获取样式ID
            style_id = tblStyle.get(qn('w:val'))
            if style_id:
                # 尝试从样式部分获取样式定义
                try:
                    styles = self.doc.styles
                    style = styles[style_id]
                    if style:
                        # 清除样式中的边框设置
                        self._clear_style_borders(style)
                except:
                    pass  # 如果无法处理样式，忽略错误
    
    def _clear_style_borders(self, style):
        """清除样式中的边框设置"""
        try:
            element = style.element
            # 查找并移除所有边框相关元素
            border_elements = element.xpath(".//w:tblBorders | .//w:tcBorders | .//w:pBdr")
            for border_elem in border_elements:
                parent = border_elem.getparent()
                if parent is not None:
                    parent.remove(border_elem)
        except:
            pass  # 如果无法处理样式，忽略错误
    
    def _clear_cell_borders(self, cell):
        """清除单个单元格的所有边框"""
        tc = cell._tc
        tcPr = tc.tcPr
        
        # 如果没有tcPr，创建一个
        if tcPr is None:
            from docx.oxml import parse_xml
            tcPr = parse_xml('<w:tcPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            tc.append(tcPr)
        
        # 直接移除所有tcBorders元素
        tcBorders_list = tcPr.xpath(".//w:tcBorders")
        for tcBorders in tcBorders_list:
            tcPr.remove(tcBorders)
        
        # 清除单元格内所有段落的边框
        for paragraph in cell.paragraphs:
            self._clear_single_paragraph_border(paragraph)
    
    def _clear_paragraph_borders(self):
        """清除所有段落的边框"""
        # 处理正文段落
        for paragraph in self.doc.paragraphs:
            self._clear_single_paragraph_border(paragraph)
        
        # 处理表格中的段落
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._clear_single_paragraph_border(paragraph)
    
    def _clear_single_paragraph_border(self, paragraph):
        """清除单个段落的边框"""
        p = paragraph._element
        pPr = p.pPr
        
        # 如果段落没有pPr，创建一个
        if pPr is None:
            from docx.oxml import parse_xml
            pPr = parse_xml('<w:pPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
            p.insert(0, pPr)
        
        # 直接移除所有pBdr元素
        pBdr_list = pPr.xpath(".//w:pBdr")
        for pBdr in pBdr_list:
            pPr.remove(pBdr)
    
    def set_font_to_microsoft_yahei(self):
        """将全文字体设置为微软雅黑，加粗字体设置为12磅（小四），第一行文字设置为13.5磅"""
        # 首先设置第一行文字的字体大小为13.5磅
        self._set_first_line_font_size()
        
        # 然后处理全文的字体
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = '微软雅黑'
                # 设置西文字体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run._element.rPr.rFonts.set(qn('w:ascii'), '微软雅黑')
                run._element.rPr.rFonts.set(qn('w:hAnsi'), '微软雅黑')
                run._element.rPr.rFonts.set(qn('w:cs'), '微软雅黑')
                
                # 设置字体大小：加粗字体设置为12磅（小四），其他字体保持原大小
                if run.bold:
                    run.font.size = Pt(12)  # 小四对应12磅
        
        # 处理表格中的文字
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '微软雅黑'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                            run._element.rPr.rFonts.set(qn('w:ascii'), '微软雅黑')
                            run._element.rPr.rFonts.set(qn('w:hAnsi'), '微软雅黑')
                            run._element.rPr.rFonts.set(qn('w:cs'), '微软雅黑')
                            
                            # 设置字体大小：加粗字体设置为12磅（小四），其他字体保持原大小
                            if run.bold:
                                run.font.size = Pt(12)  # 小四对应12磅
    
    def _set_first_line_font_size(self):
        """设置文档中第一行文字的字体大小为13.5磅"""
        # 查找文档中第一个非空段落
        first_paragraph = None
        for paragraph in self.doc.paragraphs:
            if paragraph.text.strip():  # 如果段落不为空
                first_paragraph = paragraph
                break
        
        # 如果找到了非空段落，设置其字体大小为13.5磅
        if first_paragraph:
            for run in first_paragraph.runs:
                run.font.size = Pt(13.5)
        
        # 处理表格中的第一行文字
        for table in self.doc.tables:
            # 找到第一个非空单元格
            first_cell = None
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # 如果单元格不为空
                        first_cell = cell
                        break
                if first_cell:
                    break
            
            # 如果找到了非空单元格，设置其第一个非空段落的字体大小为13.5磅
            if first_cell:
                for paragraph in first_cell.paragraphs:
                    if paragraph.text.strip():  # 如果段落不为空
                        for run in paragraph.runs:
                            run.font.size = Pt(13.5)
                        break  # 只处理第一个非空段落
    
    def clean_document_content(self):
        """清理文档内容：删除多余空格、空行、空段落，同时保留原有格式（包括加粗）"""
        # 收集需要删除的空段落
        paragraphs_to_delete = []
        
        # 处理段落
        for paragraph in self.doc.paragraphs:
            if paragraph.text.strip() == "":
                # 如果段落为空，标记为待删除
                paragraphs_to_delete.append(paragraph)
            else:
                # 如果段落不为空，清理内容
                self._clean_paragraph(paragraph)
        
        # 删除空段落
        for paragraph in paragraphs_to_delete:
            p = paragraph._element
            p.getparent().remove(p)
        
        # 处理表格中的内容
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # 收集表格中需要删除的空段落
                    cell_paragraphs_to_delete = []
                    
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip() == "":
                            # 如果段落为空，标记为待删除
                            cell_paragraphs_to_delete.append(paragraph)
                        else:
                            # 如果段落不为空，清理内容
                            self._clean_paragraph(paragraph)
                    
                    # 删除表格中的空段落
                    for paragraph in cell_paragraphs_to_delete:
                        p = paragraph._element
                        p.getparent().remove(p)
    
    def _clean_paragraph(self, paragraph):
        """清理单个段落的文本，保留原有格式"""
        # 获取段落文本
        text = paragraph.text
        
        # 删除多余的空格，但保留单词之间的单个空格
        # 先替换多个连续空格为单个空格
        text = re.sub(r' +', ' ', text)
        # 删除行首行尾空格
        text = text.strip()
        
        # 删除制表符和其他空白字符
        text = re.sub(r'[\t\u00A0\u2000-\u200A\u2028\u2029\u202F\u205F\u3000]', '', text)
        
        # 如果段落没有run，添加一个
        if not paragraph.runs:
            paragraph.add_run(text)
            return
        
        # 收集所有run的文本和格式信息
        runs_info = []
        for run in paragraph.runs:
            runs_info.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size,
                'run': run  # 保存run对象的引用
            })
        
        # 合并所有run的文本
        original_text = ''.join([info['text'] for info in runs_info])
        
        # 如果文本没有变化，不需要更新
        if original_text == text:
            return
        
        # 智能分配清理后的文本到各个run，保留格式
        self._distribute_text_to_runs(runs_info, text)
    
    def _distribute_text_to_runs(self, runs_info, cleaned_text):
        """智能分配文本到各个run，尽可能保留原有格式"""
        if not runs_info:
            return
        
        # 计算原始文本总长度
        original_total_length = sum(len(info['text']) for info in runs_info)
        cleaned_length = len(cleaned_text)
        
        if original_total_length == 0:
            # 如果原始文本为空，将所有清理后的文本放入第一个run
            runs_info[0]['run'].text = cleaned_text
            # 清空其他run的文本
            for i in range(1, len(runs_info)):
                runs_info[i]['run'].text = ""
            return
        
        # 计算每个run应该分配的字符数（按比例分配）
        current_pos = 0
        for i, info in enumerate(runs_info):
            if i == len(runs_info) - 1:
                # 最后一个run分配剩余的所有文本
                info['run'].text = cleaned_text[current_pos:]
            else:
                # 按原始文本长度比例分配
                original_ratio = len(info['text']) / original_total_length
                target_length = max(1, int(cleaned_length * original_ratio))
                
                # 确保不会超出清理后文本的长度
                end_pos = min(current_pos + target_length, cleaned_length)
                
                # 分配文本
                info['run'].text = cleaned_text[current_pos:end_pos]
                current_pos = end_pos
    
    def process_document(self):
        """处理文档：清理内容、设置字体、清除边框"""
        self.clean_document_content()
        self.set_font_to_microsoft_yahei()
        # 确保第一行文字大小为13.5磅，无论是否设置字体
        self._set_first_line_font_size()
        self.clear_table_borders()
        self._remove_all_line_elements()  # 添加这行来移除所有可能的线条元素
        self._debug_document_structure()  # 添加调试方法，输出文档结构信息
    
    def _debug_document_structure(self):
        """调试文档结构，输出可能的边框元素信息"""
        # 创建调试信息文件
        debug_info = []
        debug_info.append("=== 文档结构调试信息 ===\n")
        
        # 检查文档中的所有段落
        for i, paragraph in enumerate(self.doc.paragraphs):
            p_xml = paragraph._p.xml
            if 'border' in p_xml.lower() or 'bdr' in p_xml.lower() or 'shading' in p_xml.lower():
                debug_info.append(f"段落 {i+1} 可能包含边框或底纹元素:")
                debug_info.append(p_xml)
                debug_info.append("\n")
        
        # 检查表格中的所有段落
        for t_idx, table in enumerate(self.doc.tables):
            debug_info.append(f"\n=== 表格 {t_idx+1} ===\n")
            tbl_xml = table._tbl.xml
            if 'border' in tbl_xml.lower() or 'bdr' in tbl_xml.lower() or 'shading' in tbl_xml.lower():
                debug_info.append(f"表格 {t_idx+1} 可能包含边框或底纹元素:")
                debug_info.append(tbl_xml)
                debug_info.append("\n")
            
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    tc_xml = cell._tc.xml
                    if 'border' in tc_xml.lower() or 'bdr' in tc_xml.lower() or 'shading' in tc_xml.lower():
                        debug_info.append(f"表格 {t_idx+1} 行 {r_idx+1} 列 {c_idx+1} 可能包含边框或底纹元素:")
                        debug_info.append(tc_xml)
                        debug_info.append("\n")
                    
                    for p_idx, paragraph in enumerate(cell.paragraphs):
                        p_xml = paragraph._p.xml
                        if 'border' in p_xml.lower() or 'bdr' in p_xml.lower() or 'shading' in p_xml.lower():
                            debug_info.append(f"表格 {t_idx+1} 行 {r_idx+1} 列 {c_idx+1} 段落 {p_idx+1} 可能包含边框或底纹元素:")
                            debug_info.append(p_xml)
                            debug_info.append("\n")
        
        # 将调试信息写入文件
        with open("debug_info.txt", "w", encoding="utf-8") as f:
            f.write("\n".join(debug_info))
    
    def _remove_all_line_elements(self):
        """移除文档中所有可能的线条元素"""
        # 处理文档中的所有段落
        for paragraph in self.doc.paragraphs:
            self._remove_paragraph_line_elements(paragraph)
        
        # 处理表格中的所有段落
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._remove_paragraph_line_elements(paragraph)
        
        # 直接处理文档的XML结构，移除所有可能的边框元素
        self._remove_xml_borders()
    
    def _remove_xml_borders(self):
        """直接从XML中移除所有边框相关元素"""
        # 先处理所有表格的XML结构
        for table in self.doc.tables:
            tbl = table._tbl
            # 查找并移除表格中所有边框相关元素
            border_elements = tbl.xpath(".//w:tblBorders | .//w:tcBorders | .//w:pBdr | .//w:bdr")
            for elem in border_elements:
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)
            
            # 查找并移除表格中所有底纹元素
            shading_elements = tbl.xpath(".//w:shading")
            for elem in shading_elements:
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)
        
        # 处理所有段落的XML结构
        for paragraph in self.doc.paragraphs:
            p = paragraph._p
            # 查找并移除段落中所有边框相关元素
            border_elements = p.xpath(".//w:pBdr | .//w:bdr")
            for elem in border_elements:
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)
            
            # 查找并移除段落中所有底纹元素
            shading_elements = p.xpath(".//w:shading")
            for elem in shading_elements:
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)
    
    def _remove_paragraph_line_elements(self, paragraph):
        """移除段落中所有可能的线条元素"""
        # 检查段落属性
        pPr = paragraph._p.pPr
        if pPr is not None:
            # 移除所有可能的线条相关元素
            line_elements = pPr.xpath(".//w:pBdr | .//w:shading | .//w:top | .//w:left | .//w:bottom | .//w:right")
            for elem in line_elements:
                parent = elem.getparent()
                if parent is not None:
                    parent.remove(elem)
        
        # 检查段落中的所有run
        for run in paragraph.runs:
            rPr = run._r.rPr
            if rPr is not None:
                # 移除run中可能的线条元素
                line_elements = rPr.xpath(".//w:bdr | .//w:shading")
                for elem in line_elements:
                    parent = elem.getparent()
                    if parent is not None:
                        parent.remove(elem)
    
    def save_document(self, output_path=None):
        """
        保存文档
        :param output_path: 输出路径，如果为None则覆盖原文件
        """
        if output_path is None:
            output_path = self.file_path
        self.doc.save(output_path)
        return output_path


def process_word_document(input_file, output_file=None, clear_borders=True, set_font=True, clean_content=True):
    """
    处理Word文档的主函数
    :param input_file: 输入文件路径
    :param output_file: 输出文件路径，如果为None则覆盖原文件
    :param clear_borders: 是否清除表格边框
    :param set_font: 是否设置字体为微软雅黑
    :param clean_content: 是否清理文档内容
    :return: 处理后的文件路径
    """
    try:
        # 创建处理器实例
        processor = WordProcessor(input_file)
        
        # 执行处理操作
        if clear_borders:
            print("正在清除表格边框...")
            processor.clear_table_borders()
            # 调用额外的清除线条元素方法，确保彻底清除所有边框
            processor._remove_all_line_elements()
            processor._remove_xml_borders()
        
        if set_font:
            print("正在设置字体为微软雅黑...")
            processor.set_font_to_microsoft_yahei()
        
        if clean_content:
            print("正在清理文档内容...")
            processor.clean_document_content()
        
        # 确保第一行文字大小为13.5磅，无论是否设置字体
        print("正在设置第一行文字大小为13.5磅...")
        processor._set_first_line_font_size()
        
        # 保存文档
        output_path = processor.save_document(output_file)
        print(f"文档处理完成，已保存到: {output_path}")
        
        return output_path
    except Exception as e:
        print(f"处理文档时出错: {str(e)}")
        return None


if __name__ == "__main__":
    import sys
    
    # 检查命令行参数
    if len(sys.argv) < 2:
        print("使用方法: python word_processor.py <输入文件路径> [输出文件路径]")
        print("示例: python word_processor.py input.docx output.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(input_file):
        print(f"错误: 文件 '{input_file}' 不存在")
        sys.exit(1)
    
    # 处理文档
    result = process_word_document(input_file, output_file)
    
    if result:
        print("处理完成!")
    else:
        print("处理失败!")
        sys.exit(1)